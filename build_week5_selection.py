from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"C:\Users\josh\OneDrive\Pictures\Documents\Week 5 Project - Selection - Joshua Sabzbagh.docx")

# standard_business_brief preset
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY = RGBColor(0x55, 0x55, 0x55)
MUTED = RGBColor(0x6B, 0x72, 0x80)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
LIGHT_GRAY = "F2F4F7"
BORDER = "C9D3DF"


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=4):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError("Table widths must sum to 9360 DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell, **CELL_MARGIN_DXA)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_borders(table)


def set_bottom_border(paragraph, color="2E74B5", size=12, space=5):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    label = paragraph.add_run("Page ")
    set_run_font(label, size=9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=9, color=MUTED)


def add_numbering_definition(doc, start=1):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{0xA1000000 + abstract_id:08X}")
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    tmpl = OxmlElement("w:tmpl")
    tmpl.set(qn("w:val"), f"{0xB2000000 + abstract_id:08X}")
    abstract.append(tmpl)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start_el = OxmlElement("w:start")
    start_el.set(qn("w:val"), str(start))
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "40")
    spacing.set(qn("w:line"), "264")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    lvl.extend([start_el, num_fmt, lvl_text, suff, lvl_jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    # Force each interview-question section to begin at 1 even when Word
    # recognizes the two lists as visually similar.
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start))
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def add_body(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.widow_control = True
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_question(doc, num_id, question, rationale):
    q = doc.add_paragraph(style="Question")
    apply_numbering(q, num_id)
    q.paragraph_format.keep_with_next = True
    run = q.add_run(question)
    set_run_font(run, bold=True)

    r_p = doc.add_paragraph(style="Rationale")
    label = r_p.add_run("Rationale: ")
    set_run_font(label, bold=True, color=DARK_BLUE)
    body = r_p.add_run(rationale)
    set_run_font(body)


def add_works_cited_entry(doc, text):
    p = doc.add_paragraph(style="Works Cited Entry")
    r = p.add_run(text)
    set_run_font(r, size=10.5)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

# Exact style tokens for standard_business_brief.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = BLACK
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

for level, size, color, before, after in (
    (1, 16, BLUE, 16, 8),
    (2, 13, BLUE, 12, 6),
    (3, 12, DARK_BLUE, 8, 4),
):
    style = doc.styles[f"Heading {level}"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.keep_with_next = True

question_style = doc.styles.add_style("Question", WD_STYLE_TYPE.PARAGRAPH)
question_style.base_style = normal
question_style.font.name = "Calibri"
question_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
question_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
question_style.font.size = Pt(11)
question_style.paragraph_format.space_before = Pt(4)
question_style.paragraph_format.space_after = Pt(2)
question_style.paragraph_format.line_spacing = 1.10
question_style.paragraph_format.keep_with_next = True

rationale_style = doc.styles.add_style("Rationale", WD_STYLE_TYPE.PARAGRAPH)
rationale_style.base_style = normal
rationale_style.font.name = "Calibri"
rationale_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
rationale_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
rationale_style.font.size = Pt(11)
rationale_style.paragraph_format.left_indent = Inches(0.5)
rationale_style.paragraph_format.space_before = Pt(0)
rationale_style.paragraph_format.space_after = Pt(8)
rationale_style.paragraph_format.line_spacing = 1.10

wc_style = doc.styles.add_style("Works Cited Entry", WD_STYLE_TYPE.PARAGRAPH)
wc_style.base_style = normal
wc_style.font.name = "Calibri"
wc_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
wc_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
wc_style.font.size = Pt(10.5)
wc_style.paragraph_format.left_indent = Inches(0.5)
wc_style.paragraph_format.first_line_indent = Inches(-0.5)
wc_style.paragraph_format.space_before = Pt(0)
wc_style.paragraph_format.space_after = Pt(6)
wc_style.paragraph_format.line_spacing = 1.10

# Quiet running header.
header = section.header
header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
header_table.alignment = WD_TABLE_ALIGNMENT.LEFT
header_table.autofit = False
for idx, width in enumerate((4680, 4680)):
    set_cell_width(header_table.rows[0].cells[idx], width)
    set_cell_margins(header_table.rows[0].cells[idx], top=0, start=0, bottom=0, end=0)
header_table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
header_table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
left_run = header_table.rows[0].cells[0].paragraphs[0].add_run("WEEK 5 PROJECT")
set_run_font(left_run, size=8.5, color=MUTED, bold=True)
right_run = header_table.rows[0].cells[1].paragraphs[0].add_run("FRITO-LAY ROUTE SALES REPRESENTATIVE")
set_run_font(right_run, size=8.5, color=MUTED)

footer = section.footer
add_page_number(footer.paragraphs[0])

# memo_masthead first-page treatment.
spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(12)
title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(4)
title_run = title.add_run("WEEK 5 PROJECT — SELECTION")
set_run_font(title_run, size=23, color=BLACK, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_before = Pt(0)
subtitle.paragraph_format.space_after = Pt(16)
subtitle_run = subtitle.add_run("Frito-Lay Route Sales Representative")
set_run_font(subtitle_run, size=14, color=GRAY)

for label, value in (
    ("Student", "Joshua Sabzbagh"),
    ("Course", "MGMT2001: Human Resource Management"),
    ("Instructor", "Deet Gilbert"),
    ("Date", "16 August 2026"),
):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, bold=True)
    value_run = p.add_run(value)
    set_run_font(value_run)

rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(6)
rule.paragraph_format.space_after = Pt(10)
set_bottom_border(rule)

add_heading(doc, "Part One: Selection Process Development", 1)
add_body(
    doc,
    "The selection plan for PepsiCo’s Frito-Lay Route Sales Representative should identify a candidate who can sell, deliver, and merchandise products while managing a route safely and independently. These duties match Frito-Lay’s description of the role, including driving between stores, unloading orders, maintaining records, and working with store managers to grow sales (Frito-Lay). The process must also be consistent: every applicant should be measured against the same job-related requirements, questions, and scoring standards. A structured process improves the quality of the hiring decision and supports fairness because decisions are tied to the actual work rather than an interviewer’s personal impression (Varlaro).",
)

add_heading(doc, "A. Initial Screening Process", 2)
add_heading(doc, "Initial Candidate Selection", 3)
add_body(
    doc,
    "The process should begin when the posting closes or the target number of complete applications is reached. The HR recruiter should review the standardized application and résumé for the essential requirements identified in the earlier job analysis and the location-specific posting. These include being at least 21 years old, holding a valid driver’s license and proof of insurance, willingness to complete an acceptable motor-vehicle-record review later in the process, authorization to work in the United States, and availability for early mornings, weekends, holidays, and changing route schedules. Applicants should also receive a realistic description of frequent driving, independent field work, standing throughout the shift, and frequently lifting 40 pounds to stock and merchandise products. The employer may ask whether an applicant can perform the essential duties with or without reasonable accommodation, but it should not ask about a diagnosis, injury history, medication, or disability before a conditional offer (EEOC, “Pre-Employment Inquiries”).",
)
add_body(
    doc,
    "The recruiter should first apply pass-or-review gates to the complete application. Candidates who meet those requirements should then receive a structured preferred-qualification score. Strong evidence would include prior retail, route delivery, merchandising, customer service, sales, inventory, or consumer-products experience; accurate use of a handheld or mobile system; safe work habits; schedule reliability; and examples of managing several customer accounts or priorities. Direct Frito-Lay experience should not be required because grocery, warehouse, restaurant, delivery, and other customer-facing work may provide transferable skills. The interviewee from the earlier job analysis also emphasized organization, basic mathematics, safe driving, customer relationships, and the ability to solve routine route problems independently (Anonymous Frito-Lay Route Sales Representative).",
)

add_heading(doc, "Responsibilities and Application Review Criteria", 3)
add_body(
    doc,
    "A trained HR recruiter should conduct the initial review because HR can apply the requirements consistently and screen out non-job-related information. Before applications are reviewed, the recruiter and District Sales Leader should agree on one scoring guide and examples of acceptable evidence. The recruiter should document each decision, then send the proposed shortlist to the District Sales Leader for a calibration review. The manager may confirm whether the experience matches route needs, but may not add new preferences after seeing the candidates. This shared review gives the operation useful input while preserving consistency.",
)
add_body(
    doc,
    "An application may be removed from consideration if it is materially incomplete after a reasonable opportunity to correct it, if the applicant does not meet a stated essential requirement, or if the applicant cannot meet the disclosed schedule. HR should not reject a person for lacking a college degree because the prior job analysis did not identify a degree as necessary. HR also should not score names, photographs, age-related dates, family information, or other protected or non-job-related details. The documented rubric should focus on evidence of customer service, selling or merchandising, inventory accuracy, safe driving responsibility, time management, technology use, and dependability. The same standard should apply to internal and external candidates.",
)

add_heading(doc, "Ideal Number of Applicants", 3)
add_body(
    doc,
    "A practical target is 30 complete applications, eight first-round interviews, three second-round interviews, one contingent offer, and one documented alternate. Thirty applications should provide a useful range without making the process unnecessarily slow. If fewer than eight applicants meet the minimum requirements, PepsiCo should extend or broaden recruiting instead of lowering the standard. If many more than eight qualify, the preferred-qualification rubric should determine the shortlist, with HR reviewing any tied scores.",
)

funnel_table = doc.add_table(rows=1, cols=3)
funnel_table.style = "Table Grid"
headers = ("Selection stage", "Target", "Decision rule")
for i, text in enumerate(headers):
    cell = funnel_table.rows[0].cells[i]
    shade_cell(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=10, bold=True, color=DARK_BLUE)
set_repeat_table_header(funnel_table.rows[0])
for stage, target, rule_text in (
    ("Complete applications", "30", "All required application fields and pass-or-review gates"),
    ("First-round interviews", "8", "Highest standardized screening scores"),
    ("Second-round panel", "3", "Meets the first-round minimum and ranks in the top three"),
    ("Contingent offer", "1 + alternate", "Highest documented total score; alternate retained if needed"),
):
    row = funnel_table.add_row()
    for i, value in enumerate((stage, target, rule_text)):
        p = row.cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(value)
        set_run_font(r, size=10)
set_table_geometry(funnel_table, [2700, 1500, 5160])

add_heading(doc, "B. Interview Process", 2)
add_heading(doc, "First-Round Interview and Work Sample", 3)
add_body(
    doc,
    "The first round should be a 35-minute structured video or telephone interview conducted by the HR recruiter and the District Sales Leader. The same five situational questions should be asked in the same order. The interviewers should use a one-to-five anchored scale that describes weak, acceptable, and strong evidence for customer service, safety judgment, time management, inventory decisions, and sales communication. Each interviewer should score independently before the scores are discussed. This approach makes the comparison more reliable and keeps the discussion centered on the requirements of the route.",
)
add_body(
    doc,
    "Candidates should then complete a short job sample rather than a broad intelligence, personality, or physical test. The sample should provide a mock route with store delivery windows, simple inventory and sales figures, a missing product, and one customer concern. The candidate would explain the order of stops, calculate a reasonable order adjustment, and describe the customer communication needed. This exercise directly measures basic numerical reasoning, organization, attention to inventory, and practical judgment. The EEOC advises employers to ensure that selection procedures are job-related, properly validated, accessible when accommodation is needed, and reviewed for adverse impact (EEOC, “Employment Tests”).",
)
add_body(
    doc,
    "No medical examination or general physical-capacity test should occur before the offer. At the interview stage, the employer should explain the essential tasks and ask whether the applicant can perform them with or without reasonable accommodation. The first-round score should combine the structured interview (60 percent) and the route work sample (40 percent). Candidates must earn at least 60 percent on the work sample and 70 percent on the first-round composite. The three highest-scoring candidates who meet both standards should advance. Using both a minimum standard and rank prevents an unqualified candidate from advancing merely because the pool is small.",
)

add_heading(doc, "Second-Round Interview", 3)
add_body(
    doc,
    "The second round should be an in-person panel interview with the District Sales Leader, a trained experienced Route Sales Representative, and an HR representative. The candidate should first receive a realistic preview of the route, schedule variability, customer interaction, sales expectations, recordkeeping, and physical work. The panel should then ask the same five behavioral questions about actual past experiences. These questions should cover difficult customers, managing several priorities, ordering or inventory decisions, selling a product or display idea, and completing delivery or stocking work safely. The candidate should also have time to ask questions so that both sides can judge whether the role is a realistic fit.",
)
add_body(
    doc,
    "Panel members should score responses independently with the same one-to-five anchors before a group discussion. The final 100-point decision should include application qualifications (15 points), the first-round structured interview (20 points), the route work sample (20 points), and the second-round behavioral interview (45 points). HR should audit the score sheet for missing ratings, unsupported comments, or inconsistent standards. The panel should identify the highest-scoring candidate and one alternate, record the job-related reasons for the decision, and retain the documentation according to company policy.",
)

add_heading(doc, "C. Contingent Job Offer", 2)
add_body(
    doc,
    "The selected candidate should receive a written contingent offer stating the job title, territory or reporting location, starting pay, expected schedule, benefits eligibility, supervisor, proposed start date, deadline to respond, and any applicable at-will language. The letter should list every condition clearly and explain that employment will not begin until HR confirms that the conditions are satisfied. The offer should not promise a fixed route or schedule unless PepsiCo intends to guarantee it.",
)
add_body(
    doc,
    "Because driving is an essential duty, the first condition should be verification of the driver’s license and an acceptable motor vehicle record under a standard used for all Route Sales Representative candidates. This review is directly related to roadway safety, company property, customers, and insurance requirements. If the specific vehicle or route is covered by Department of Transportation rules, HR should also complete the applicable DOT qualifications. The plan should not assume that every route is DOT-regulated; HR should apply only the requirements that match the assigned vehicle and route.",
)
add_body(
    doc,
    "With the candidate’s written authorization, HR may verify employment and job-related references and obtain a background report. The FTC explains that when a third-party background company is used, the employer must provide a stand-alone disclosure, obtain written permission, and follow the required pre-adverse and adverse-action process if report information affects the decision (FTC). A criminal record should not produce an automatic rejection. HR should consider the nature and seriousness of the conduct, the time that has passed, and its relationship to driving, customer locations, products, equipment, or other actual duties (EEOC, “Arrest and Conviction Records”). A credit check should not be used because the job analysis did not establish it as necessary.",
)
add_body(
    doc,
    "After the conditional offer, PepsiCo may use a job-related functional assessment that reflects the actual requirement to frequently handle 40-pound cases and the related carrying, pushing, stocking, bending, and standing demands, provided it is administered consistently to all entering employees in the same job category and reasonable accommodation is considered. The EEOC permits post-offer medical questions or examinations when the same requirement is applied to everyone entering that job category; medical information must remain confidential (EEOC, “Pre-Employment Inquiries”). If PepsiCo’s written policy, the assigned vehicle classification, and applicable law require a drug or alcohol screen, it should also occur after the offer and be administered uniformly. Its purpose should be route and workplace safety, not an unrelated inquiry into private conduct.",
)
add_body(
    doc,
    "Finally, HR should complete identity and work-authorization verification at the legally appropriate time after acceptance. Once the motor vehicle review, reference and employment checks, background process, any applicable post-offer assessment or substance screen, and required work-authorization documents are complete, HR can confirm the final offer and schedule orientation. If a condition is not satisfied, HR should document the job-related reason, provide any notices or opportunity to correct inaccurate information required by law, and then consider the alternate candidate. This sequence protects safety and business needs while treating applicants consistently and respectfully.",
)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
add_heading(doc, "Part Two: Structured Interview Development", 1)
add_body(
    doc,
    "The following questions are tied to the essential duties identified in the job analysis. Interviewers should ask each candidate the same core questions and use the same anchored scoring guide. Behavioral questions request an actual past example; situational questions ask how the applicant would respond to a realistic future event. None of the questions requests protected, medical, or family information.",
)

behavioral_num = add_numbering_definition(doc)
# Continue the same real numbered list across the two clearly labeled groups.
# This produces Behavioral 1–5 and Situational 6–10 without manual numbering.
situational_num = behavioral_num

add_heading(doc, "Behavioral Interview Questions", 2)
for question, rationale in (
    (
        "Tell us about a time you handled a dissatisfied customer or store manager who raised a concern about service, product availability, or merchandise placement. What did you do, and what was the outcome?",
        "This evaluates conflict resolution, communication, and relationship-building skills, which are essential for maintaining retail accounts and protecting customer satisfaction.",
    ),
    (
        "Describe a time you had several deliveries, customer commitments, or work priorities to complete within a limited period. How did you organize your schedule and respond to unexpected delays?",
        "This assesses route planning, time management, and adaptability—key requirements for completing an assigned route accurately and on schedule.",
    ),
    (
        "Give an example of a time you used sales, inventory, or customer-demand information to make an ordering decision. How did you avoid shortages, excess inventory, or product waste?",
        "This measures judgment and attention to inventory levels, which help a Route Sales Representative keep shelves stocked while limiting returns, stale product, and unnecessary waste.",
    ),
    (
        "Tell us about a time you persuaded a customer to support a new product, promotion, display, or sales idea. How did you identify the opportunity and present its value?",
        "This evaluates selling ability, initiative, and customer-focused communication, all of which support product growth and successful promotion execution.",
    ),
    (
        "Describe a past delivery or stocking assignment that required you to move and merchandise products while following safety and quality procedures. How did you complete the work safely and accurately?",
        "This assesses safe work habits, attention to detail, and experience performing essential delivery and merchandising tasks without asking about medical conditions or disabilities.",
    ),
):
    add_question(doc, behavioral_num, question, rationale)

add_heading(doc, "Situational Interview Questions", 2)
for question, rationale in (
    (
        "Suppose a store manager is upset because a popular Frito-Lay product was missing from the delivery, but you still have a full route to complete. How would you respond and resolve the issue?",
        "This evaluates customer service, accountability, and practical problem solving while balancing one account’s concern with the needs of the remaining route.",
    ),
    (
        "Imagine that traffic and a vehicle delay put you behind schedule, and several stores have preferred delivery windows. How would you prioritize the remaining stops and communicate with customers?",
        "This assesses route management, decision making, and proactive communication under time pressure.",
    ),
    (
        "At a store, shelf inventory is low, but demand has recently been inconsistent and the backroom count may be inaccurate. How would you decide what quantity to order and deliver?",
        "This measures inventory judgment, use of available information, and the ability to balance product availability against overstock and waste.",
    ),
    (
        "Suppose Frito-Lay launches a promotion that could increase sales, but a store manager is reluctant to provide display space. How would you present the opportunity and respond to the manager’s concerns?",
        "This evaluates consultative selling, product communication, negotiation, and the ability to build support without damaging the customer relationship.",
    ),
    (
        "Imagine that you arrive at a stop and find the unloading area obstructed, making the planned delivery method unsafe. What steps would you take to protect people and products while completing the delivery appropriately?",
        "This assesses safety judgment, compliance with procedures, and problem solving during the physical delivery process.",
    ),
):
    add_question(doc, situational_num, question, rationale)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
add_heading(doc, "Works Cited", 1)
add_works_cited_entry(
    doc,
    "Anonymous Frito-Lay Route Sales Representative. Personal interview. 26 July 2026.",
)
add_works_cited_entry(
    doc,
    "Equal Employment Opportunity Commission. “Employment Tests and Selection Procedures.” U.S. Equal Employment Opportunity Commission, https://www.eeoc.gov/laws/guidance/employment-tests-and-selection-procedures. Accessed 16 Aug. 2026.",
)
add_works_cited_entry(
    doc,
    "Equal Employment Opportunity Commission. “Enforcement Guidance on the Consideration of Arrest and Conviction Records in Employment Decisions under Title VII of the Civil Rights Act.” U.S. Equal Employment Opportunity Commission, 25 Apr. 2012, https://www.eeoc.gov/laws/guidance/enforcement-guidance-consideration-arrest-and-conviction-records-employment-decisions. Accessed 16 Aug. 2026.",
)
add_works_cited_entry(
    doc,
    "Equal Employment Opportunity Commission. “Pre-Employment Inquiries and Medical Questions & Examinations.” U.S. Equal Employment Opportunity Commission, https://www.eeoc.gov/pre-employment-inquiries-and-medical-questions-examinations. Accessed 16 Aug. 2026.",
)
add_works_cited_entry(
    doc,
    "Federal Trade Commission. “Background Checks: What Employers Need to Know.” Federal Trade Commission, 11 Mar. 2014, https://www.ftc.gov/business-guidance/resources/background-checks-what-employers-need-know. Accessed 16 Aug. 2026.",
)
add_works_cited_entry(
    doc,
    "Frito-Lay. “Route Sales Representative (Centennial, Colorado).” Frito-Lay Employment, PepsiCo, https://www.fritolayemployment.com/jobs/route-sales-representative-centennial-colorado-480303. Accessed 16 Aug. 2026.",
)
add_works_cited_entry(
    doc,
    "Varlaro, John D. The Practice of Human Resource Management. 2nd ed., Great River Learning, 2022.",
)

# Core properties and save.
doc.core_properties.title = "Week 5 Project — Selection: Frito-Lay Route Sales Representative"
doc.core_properties.subject = "MGMT2001 Human Resource Management"
doc.core_properties.author = "Joshua Sabzbagh"
doc.core_properties.keywords = "selection process, structured interview, route sales representative, HRM"

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
